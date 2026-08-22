from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import json


def document_generator(data_json: str) -> str:
    # Convert JSON string into Python dictionary
    data = json.loads(data_json)

    # Create document
    document = Document()

    # =========================
    # Title
    # =========================
    title = document.add_heading("User Problem Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_paragraph = document.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_run = date_paragraph.add_run(
        datetime.now().strftime("%B %d, %Y")
    )
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(100, 100, 100)

    # =========================
    # Personal Information
    # =========================
    document.add_heading("Personal Information", level=1)

    user_data = data["user_data"]

    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = user_data["name"]

    table.cell(1, 0).text = "Email"
    table.cell(1, 1).text = user_data["email"]

    # =========================
    # Problem
    # =========================
    document.add_heading("Problem", level=1)

    problem_paragraph = document.add_paragraph()
    problem_paragraph.add_run(data["problem"])

    # =========================
    # Tips
    # =========================
    document.add_heading("Recommendations", level=1)

    tips = data["tips"].split(",")

    for tip in tips:
        tip = tip.strip()

        if tip:
            document.add_paragraph(
                tip,
                style="List Bullet"
            )

    # =========================
    # Footer
    # =========================
    section = document.sections[0]
    footer = section.footer

    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer_run = footer_paragraph.add_run(
        "Generated automatically"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(120, 120, 120)

    # TODO: Add artifact method and maybe load to gdrive
    # Save document
    filename = "user_problem_report.docx"
    document.save(filename)

    return filename